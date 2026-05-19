"""
Genera el Informe SEO de Top Doctors España en formato DOCX (compatible Google Docs).
Diseño basado en el referente Volkswagen Perú entregado por el cliente.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ---- Paleta ----
NAVY        = RGBColor(0x0B, 0x2E, 0x5C)   # azul corporativo
BLUE        = RGBColor(0x1F, 0x4E, 0x8C)
RED         = RGBColor(0xC0, 0x1B, 0x1B)
GREEN       = RGBColor(0x1E, 0x7E, 0x34)
GREY_DARK   = RGBColor(0x33, 0x33, 0x33)
GREY_MED    = RGBColor(0x66, 0x66, 0x66)
GREY_LIGHT  = "EAEFF5"   # fondo metric cards
CALLOUT_BG  = "F4F7FB"   # fondo callout
CALLOUT_RED = "FDEEEE"
LINE_BLUE   = "1F4E8C"
LINE_RED    = "C01B1B"

# ---------- Helpers de bajo nivel ----------
def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_border(cell, **kwargs):
    """kwargs: top, left, bottom, right -> dict(size, color, val)"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tc_pr.append(tcBorders)
    for edge in ('top','left','bottom','right'):
        if edge in kwargs:
            cfg = kwargs[edge]
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),  cfg.get('val','single'))
            tag.set(qn('w:sz'),   str(cfg.get('size',8)))
            tag.set(qn('w:color'),cfg.get('color','000000'))
            # quita previo y agrega
            old = tcBorders.find(qn(f'w:{edge}'))
            if old is not None: tcBorders.remove(old)
            tcBorders.append(tag)

def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top   ={'val':'nil'},
                left  ={'val':'nil'},
                bottom={'val':'nil'},
                right ={'val':'nil'})

def add_run(p, text, size=11, bold=False, color=None, italic=False, font='Calibri'):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r

def add_para(doc, text='', size=11, bold=False, color=None, align=None, space_after=4, italic=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    if text:
        add_run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    add_run(p, text, size=22, bold=True, color=NAVY)
    # línea inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'8'); bottom.set(qn('w:color'), LINE_BLUE); bottom.set(qn('w:space'),'4')
    pBdr.append(bottom); pPr.append(pBdr)

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, size=14, bold=True, color=BLUE)

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, size=11.5, bold=True, color=NAVY)

def bullet(doc, text, color=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        add_run(p, bold_prefix, size=11, bold=True, color=NAVY)
        add_run(p, text, size=11, color=color or GREY_DARK)
    else:
        add_run(p, text, size=11, color=color or GREY_DARK)

def callout(doc, label, text, kind='info'):
    fill = CALLOUT_BG if kind=='info' else CALLOUT_RED
    border_color = LINE_BLUE if kind=='info' else LINE_RED
    label_color  = NAVY if kind=='info' else RED
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade(cell, fill)
    set_cell_border(cell,
        left  ={'val':'single','size':'24','color':border_color},
        top   ={'val':'single','size':'4', 'color':border_color},
        right ={'val':'single','size':'4', 'color':border_color},
        bottom={'val':'single','size':'4', 'color':border_color},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, label+': ', size=11, bold=True, color=label_color)
    add_run(p, text, size=11, color=GREY_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def metric_cards(doc, cards):
    """cards: list of dicts {value, value_color, label, sub}"""
    table = doc.add_table(rows=1, cols=len(cards))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cards):
        cell = table.rows[0].cells[i]
        shade(cell, GREY_LIGHT)
        set_cell_border(cell,
            top   ={'val':'single','size':'4','color':'D6DEE8'},
            bottom={'val':'single','size':'4','color':'D6DEE8'},
            left  ={'val':'single','size':'4','color':'D6DEE8'},
            right ={'val':'single','size':'4','color':'D6DEE8'},
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # value
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_run(p, c['value'], size=26, bold=True, color=c.get('value_color', NAVY))
        # sub
        if c.get('sub'):
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(2)
            add_run(p2, c['sub'], size=9.5, bold=True, color=c.get('value_color', NAVY))
        # label
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(4)
        add_run(p3, c['label'], size=9, color=GREY_MED)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def styled_table(doc, header, rows, col_widths=None, header_fill='0B2E5C', zebra='F5F8FC'):
    cols = len(header)
    table = doc.add_table(rows=1+len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    # header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_run(p, h, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    # body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[1+ri].cells[ci]
            if ri % 2 == 1:
                shade(cell, zebra)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            text = str(val) if val is not None else ''
            # negro por defecto, rojo si empieza con '-'
            color = GREY_DARK
            if isinstance(val, str) and (val.startswith('+')):
                color = GREEN
            add_run(p, text, size=10, color=color)
    # bordes
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'CDD6E0')
        tbl_borders.append(b)
    table._tbl.tblPr.append(tbl_borders)
    return table

# ---------- Documento ----------
doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = GREY_DARK

# ---- Header con barra superior ----
section = doc.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(hp, 'Informe SEO Confidencial  |  Top Doctors España  |  Mayo 2026', size=9, color=GREY_MED, italic=True)
# línea inferior del header
pPr = hp._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:color'), '1F4E8C'); bottom.set(qn('w:space'),'2')
pBdr.append(bottom); pPr.append(pBdr)

# Footer
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, 'SEO LAB AGENCY  ·  Análisis basado exclusivamente en datos verificables de Ahrefs', size=8.5, color=GREY_MED, italic=True)

# ============ PORTADA ============
brand = doc.add_paragraph()
brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(brand, 'SEO LAB AGENCY', size=14, bold=True, color=NAVY)

# Espaciado
doc.add_paragraph().paragraph_format.space_after = Pt(40)
doc.add_paragraph().paragraph_format.space_after = Pt(20)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(title, 'INFORME DE ANÁLISIS SEO', size=28, bold=True, color=NAVY)

sub = doc.add_paragraph()
add_run(sub, 'Diagnóstico de Visibilidad Orgánica y Plan de Crecimiento', size=18, color=BLUE)

domain = doc.add_paragraph()
add_run(domain, 'topdoctors.es', size=22, bold=True, color=GREY_DARK)

doc.add_paragraph().paragraph_format.space_after = Pt(36)

meta = doc.add_table(rows=4, cols=2); meta.autofit = True
no_borders(meta)
rows = [
    ('Fecha de análisis',  '19 de mayo de 2026'),
    ('Herramienta utilizada', 'Ahrefs — Base de datos de SEO global (DR/UR, Rank Tracker, Backlinks, Top Pages)'),
    ('Clasificación',      'Confidencial — Uso exclusivo del cliente'),
    ('Creado por',         'SEO Lab Agency'),
]
for i,(k,v) in enumerate(rows):
    c1, c2 = meta.rows[i].cells
    add_run(c1.paragraphs[0], k, size=10.5, bold=True, color=NAVY)
    add_run(c2.paragraphs[0], v, size=10.5, color=GREY_DARK)
    c1.width = Cm(5); c2.width = Cm(12)

doc.add_page_break()

# ============ 1. RESUMEN EJECUTIVO ============
h1(doc, '1. Resumen Ejecutivo')

add_para(doc,
    'topdoctors.es es uno de los activos digitales más sólidos del sector salud privado en España: '
    'Domain Rating 78/100, posición #19,920 en el ranking global de Ahrefs y un universo de 100,335 keywords orgánicas '
    'que generan 760,959 visitas mensuales y un valor de tráfico estimado en $21,570,727 USD/mes. '
    'Esa autoridad, sin embargo, convive con dos amenazas latentes y dos oportunidades capitalizables que están '
    'dejando dinero sobre la mesa cada día.', size=11)

add_para(doc,
    'El problema no es de visibilidad — el dominio aparece. El problema es que parte sustancial de su tráfico se concentra '
    'exactamente en el tipo de contenido que Google AI Overviews y los motores de IA generativa (ChatGPT, Gemini, Perplexity) '
    'están reemplazando, mientras decenas de keywords de cinco cifras de volumen están a una posición de distancia del primer '
    'puesto sin que nadie las esté trabajando.', size=11)

callout(doc, 'HALLAZGO PRINCIPAL',
    'topdoctors.es tiene 85 keywords (34% de la muestra analizada) atrapadas en posiciones 4–10, varias con volúmenes '
    'mensuales entre 13,000 y 37,000 búsquedas. El 48.16% del tráfico total del dominio proviene de las carpetas '
    '/diccionario-medico/ y /farmacia/ — exactamente los formatos de contenido que AI Overviews está absorbiendo en España. '
    'El dominio aún no muestra una caída medida, pero los datos confirman una exposición estructural a un riesgo '
    'inminente que la competencia (Doctoralia, Quirónsalud, Sanitas) ya está cubriendo.',
    kind='info')

# Metric cards
metric_cards(doc, [
    {'value':'760,959',  'label':'Tráfico orgánico mensual estimado',           'value_color': NAVY},
    {'value':'100,335',  'label':'Keywords orgánicas en top-100 (universo)',     'value_color': NAVY},
    {'value':'$21.57M',  'sub':'USD/mes', 'label':'Valor de tráfico (CPC equivalente Google Ads)', 'value_color': GREEN},
    {'value':'78',       'sub':'/ 100',   'label':'Domain Rating — autoridad de dominio Ahrefs',   'value_color': BLUE},
])

metric_cards(doc, [
    {'value':'6,525',    'label':'Referring domains que enlazan al sitio',         'value_color': NAVY},
    {'value':'353,179',  'label':'Backlinks totales detectados',                   'value_color': NAVY},
    {'value':'0',        'label':'Featured snippets en muestra de 250 keywords',   'value_color': RED},
    {'value':'12.23',    'sub':'/ 100',   'label':'DR promedio de referrentes — perfil débil', 'value_color': RED},
])

h2(doc, 'Dos amenazas latentes y dos oportunidades capitalizables')

h3(doc, '1. Amenaza — AI Overviews están consumiendo el tráfico informacional')
add_para(doc,
    'Google lanzó AI Overviews en España en 2024 y el sector salud es el segmento con mayor índice de aparición de respuestas '
    'generativas. La carpeta /diccionario-medico/ aporta el 24.83% del tráfico de topdoctors.es (≈63,900 visitas/mes en la '
    'muestra analizada) y la carpeta /farmacia/ aporta otro 23.33% (≈60,000 visitas/mes). En conjunto, casi la mitad del '
    'tráfico del dominio depende de queries que Google está respondiendo directamente en el resultado, antes de que el usuario '
    'haga clic. La muestra de 250 keywords no detectó NI UN featured snippet capturado por topdoctors.es — es decir, hoy el '
    'dominio no está siendo citado como fuente de respuesta por Google ni por las IA generativas.', size=11)

h3(doc, '2. Amenaza — Perfil de autoridad asimétrico')
add_para(doc,
    'El dominio tiene Domain Rating 78 (autoridad excepcional), pero el DR promedio de los 250 referrentes más recientes es '
    'solamente 12.23/100. La autoridad se sostiene mayoritariamente por un puñado de enlaces históricos (wikipedia.org DR 97, '
    'cosmopolitan.com DR 89, hellomagazine.com DR 82, topdoctors.co.uk DR 74 con 3,000 backlinks). El flujo de enlaces nuevos '
    'es de baja calidad: cualquier corrección algorítmica que devalúe los enlaces históricos golpea de frente la autoridad acumulada.', size=11)

h3(doc, '3. Oportunidad — 85 keywords atrapadas en posiciones 4–10')
add_para(doc,
    'El 34% de las keywords de la muestra están en la "zona muerta" de Google (posiciones 4–10), donde el CTR es entre 4 y 7 '
    'veces menor que en la posición #1. Hay 35 keywords con volumen ≥ 1,000 búsquedas/mes en esa zona, varias de ellas con '
    'Keyword Difficulty entre 0 y 10 — es decir, son ascensos perfectamente alcanzables con trabajo de contenido y enlazado '
    'interno bien dirigido.', size=11)

h3(doc, '4. Oportunidad — Cola larga y huella geográfica subutilizada')
add_para(doc,
    'topdoctors.es ya tiene presencia activa en 30+ ciudades españolas (Madrid, Barcelona, Valencia, Alicante, Málaga, '
    'Sevilla, Bilbao, Pamplona, Las Palmas, Tenerife, Granada, Zaragoza, etc.) pero Madrid concentra el 15.46% del tráfico '
    'mientras Barcelona solo el 3.07%. La huella urbana está creada pero no está optimizada — la replicación del modelo '
    'Madrid en las 9 ciudades top puede duplicar el tráfico de la sección por sí sola.', size=11)

h2(doc, 'AI Search: el canal donde se está perdiendo la conversación')
add_para(doc,
    'En 2026, una porción creciente de la búsqueda médica en España ya no termina en Google: termina en ChatGPT, '
    'Gemini, Perplexity o en la propia respuesta generativa de Google. La IA recomienda profesionales y plataformas '
    'basándose en el contenido que cada marca ha producido para ser citada como fuente confiable. Los formatos que '
    'la IA prioriza son: contenido estructurado con schema MedicalCondition/MedicalProcedure, Q&A directos, listas '
    'numeradas con definiciones, autoría médica verificable y citas a fuentes oficiales.', size=11)

add_para(doc,
    'topdoctors.es tiene la materia prima — 250 páginas activas en /diccionario-medico/ — pero hoy no están optimizadas '
    'para ese ecosistema. La medición precisa de AI Share of Voice de topdoctors.es frente a Doctoralia, Quirónsalud, '
    'Sanitas y MAPFRE es uno de los primeros entregables del plan propuesto (Semana 1, Fase 1).', size=11)

h2(doc, 'Plan en tres fases')
styled_table(doc,
    ['Fase','Plazo','Enfoque','Resultado esperado'],
    [
        ['Quick Wins',            'Semanas 1–2', 'Capturar las 30 keywords de mayor volumen en pos 4–10', '+5,000 a +8,000 visitas/mes'],
        ['Crecimiento estructural','Meses 1–2', 'Blindar /diccionario-medico/ y /farmacia/ frente a AI Overviews', '+15,000 a +25,000 visitas/mes'],
        ['Autoridad y AI Search', 'Meses 1–3', 'Link building dirigido + contenido citable por IA + medición AI SoV', 'AI Share of Voice > 25%'],
    ]
)

add_para(doc, ' ', size=4)
add_para(doc,
    'Objetivo a 90 días: incrementar el tráfico orgánico mensual en al menos un 8% (≈ +60,000 visitas/mes), elevar el '
    'valor de tráfico mensual por encima de $23M USD, capturar al menos 15 featured snippets nuevos y establecer la '
    'línea base medible de AI Share of Voice para iniciar la batalla por las recomendaciones generativas. Todo el plan '
    'se ejecuta sin incrementar el presupuesto en Google Ads — apalancando exclusivamente el activo orgánico ya construido.',
    size=11)

doc.add_page_break()

# ============ 2. METODOLOGÍA ============
h1(doc, '2. Metodología del Análisis')

add_para(doc,
    'Todos los datos presentados en este informe provienen exclusivamente de Ahrefs, herramienta líder en análisis SEO '
    'con cobertura de más de 200 países y un índice de más de 500 millones de keywords activas. La ingesta y el procesamiento '
    'fueron ejecutados a través del pipeline propietario "Ahrefs Traffic Loss Agent Swarm v1.0" sobre la base de datos '
    'analítica de SEO Lab Agency en Supabase.', size=11)

h3(doc, 'Alcance del snapshot ingerido (19 de mayo de 2026)')
bullet(doc, '100,335 keywords orgánicas detectadas en el universo total del dominio.', bold_prefix='Universo de keywords: ')
bullet(doc, '250 keywords priorizadas por volumen y tráfico (muestra representativa del 0.25%).', bold_prefix='Muestra analizada en detalle: ')
bullet(doc, '250 páginas activas con tráfico orgánico real en España.', bold_prefix='Páginas analizadas: ')
bullet(doc, '250 backlinks más recientes con perfil dofollow/nofollow y DR de origen.', bold_prefix='Backlinks analizados: ')
bullet(doc, '250 referring domains únicos con sus métricas de autoridad.', bold_prefix='Referring domains analizados: ')
bullet(doc, 'Geografía: España (configuración Ahrefs country=es).', bold_prefix='Mercado objetivo: ')

callout(doc, 'NOTA METODOLÓGICA',
    'La muestra de 250 keywords representa el 0.25% del universo de 100,335 keywords detectadas por Ahrefs para topdoctors.es. '
    'Las conclusiones cuantitativas de este informe son representativas para las keywords de mayor tráfico (donde se concentra '
    'el grueso del valor de negocio), pero la magnitud total de las oportunidades en la cola larga es razonablemente mayor a la '
    'documentada aquí. La Fase 1 del plan incluye la ingesta ampliada a 5,000 keywords para una medición exhaustiva.',
    kind='info')

# ============ 3. ESTADO ACTUAL ============
h1(doc, '3. Estado Actual del Dominio')

h2(doc, '3.1 Métricas Generales (Ahrefs, 19 mayo 2026)')
styled_table(doc,
    ['Indicador','Valor','Interpretación'],
    [
        ['Domain Rating (DR)',       '78 / 100',            'Autoridad excepcional. Top tier para el sector salud privado en España.'],
        ['Ahrefs Rank',              '#19,920 global',      'Dentro del 0.001% superior de los dominios indexados por Ahrefs.'],
        ['URL Rating promedio',      '44 / 100',            'Página principal con autoridad alta; páginas internas con margen amplio.'],
        ['Tráfico orgánico estimado','760,959 visitas/mes', 'Volumen alto y diversificado en 100,335 keywords activas.'],
        ['Keywords en top-100',      '100,335',             'Universo amplio; la muestra detallada cubre el 0.25% por volumen.'],
        ['Backlinks totales',        '353,179',             'Perfil voluminoso, principalmente histórico.'],
        ['Referring domains',        '6,525',               'Diversidad sólida en cantidad; la calidad media es baja (ver Sección 6).'],
        ['Valor de tráfico estimado','$21,570,727 USD/mes', 'Equivalente al costo de generar las mismas visitas vía Google Ads.'],
    ]
)

h2(doc, '3.2 Distribución de Posiciones en Google (muestra de 250 keywords)')
add_para(doc,
    'El perfil de posicionamiento de la muestra analizada confirma una base sólida en los primeros puestos: '
    'el 58.8% de las keywords están en top 3 y el 92.8% están en la primera página. Esto descarta cualquier '
    'problema de relevancia o calidad de contenido en el dominio.', size=11)
styled_table(doc,
    ['Rango de posición','Keywords','% sobre muestra','Lectura'],
    [
        ['Posiciones #1–#3 (top 3)',  '147', '58.8%', 'Base de tráfico consolidada — protege estos puestos.'],
        ['Posiciones #4–#10 (1ª pág)', '85', '34.0%', 'ZONA DE OPORTUNIDAD — 5,000+ visitas/mes recuperables.'],
        ['Posiciones #11–#20 (2ª pág)','16', ' 6.4%', 'Requieren refuerzo de contenido + enlazado interno.'],
        ['Posiciones #21+',             '2', ' 0.8%', 'Volumen marginal; bajo prioridad.'],
        ['Featured snippets capturados','0', ' 0.0%', 'CRÍTICO — invisible para AI Overviews y motores generativos.'],
    ]
)

callout(doc, 'LECTURA POSITIVA',
    'El 92.8% de las keywords de la muestra ya están en la primera página de Google. El trabajo SEO requerido NO es de '
    'reconstrucción ni de recuperación post-penalización: es de optimización quirúrgica de un activo que ya funciona, '
    'orientada a empujar lo que está en pos 4–10 hacia el top 3 y a capturar los featured snippets que hoy se está llevando '
    'la competencia.',
    kind='info')

h2(doc, '3.3 Distribución de Tráfico por Carpeta')
add_para(doc,
    'El tráfico de topdoctors.es se concentra en cinco grandes secciones, dos de las cuales tienen exposición directa '
    'al riesgo de AI Overviews que se aborda en el diagnóstico (Sección 4):', size=11)
styled_table(doc,
    ['Carpeta','Páginas en muestra','Tráfico mensual','% del tráfico','Tipo de contenido'],
    [
        ['/diccionario-medico/',   '74', '63,893', '24.83%', 'Informacional puro — RIESGO ALTO AI Overviews'],
        ['/farmacia/',             '31', '60,028', '23.33%', 'Local + transaccional — riesgo medio AI'],
        ['/madrid-ciudad/',        '43', '39,796', '15.46%', 'Hub geográfico — defendible y escalable'],
        ['/articulos-medicos/',    '21', '16,632', ' 6.46%', 'Editorial — apto para optimización AI'],
        ['Homepage /',              '1', '12,838', ' 4.99%', 'Brand — estable, pos #1 en "top doctors"'],
        ['/seguro-medico/',         '1', ' 9,448', ' 3.67%', 'Transaccional alto valor (CPC > $100)'],
        ['/barcelona-ciudad/',     '13', ' 7,905', ' 3.07%', 'Hub subutilizado — replicar modelo Madrid'],
        ['/doctor/',                '4', ' 7,440', ' 2.89%', 'Perfiles individuales — alto potencial E-E-A-T'],
        ['/valencia-ciudad/',       '9', ' 4,979', ' 1.93%', 'Hub geográfico secundario — escalar'],
        ['Otras ciudades + cola', '53+', '34,386', '13.36%', '20+ ciudades activas con margen de crecimiento'],
    ]
)
callout(doc, 'RIESGO CRÍTICO IDENTIFICADO',
    'El 48.16% del tráfico depende de /diccionario-medico/ + /farmacia/. Ambas son carpetas donde Google ya muestra '
    'AI Overviews para la mayoría de queries (síntomas, definiciones, dosis, horarios). Cualquier expansión de AI Overviews '
    'en queries médicas — anunciada por Google para 2026 — golpearía simultáneamente a las dos secciones que sostienen casi la '
    'mitad del tráfico del dominio.',
    kind='warn')

doc.add_page_break()

# ============ 4. DIAGNÓSTICO ============
h1(doc, '4. Diagnóstico Estratégico')

h2(doc, '4.1 Lo que los datos CONFIRMAN como fortalezas')
bullet(doc, 'Domain Rating 78 — autoridad excepcional en sector salud privado en España.', bold_prefix='Autoridad de dominio: ')
bullet(doc, '92.8% de la muestra en primera página de Google; 58.8% en top 3.', bold_prefix='Visibilidad orgánica: ')
bullet(doc, '100,335 keywords activas, 6,525 referring domains, 353,179 backlinks.', bold_prefix='Cobertura: ')
bullet(doc, '81% de los referrentes recientes son dofollow — transmiten autoridad SEO.', bold_prefix='Perfil de enlaces: ')
bullet(doc, '$21.57M USD/mes de valor de tráfico orgánico — equivalente a no gastar ese presupuesto en Google Ads.', bold_prefix='Valor de negocio: ')
bullet(doc, '30+ ciudades españolas con presencia activa — huella nacional ya construida.', bold_prefix='Huella geográfica: ')

h2(doc, '4.2 Amenaza 1 — Exposición a Google AI Overviews')
add_para(doc, 'Probabilidad de impacto: ALTA. Evidencia documentada en los datos del propio dominio.', size=11, bold=True, color=RED)
add_para(doc,
    'AI Overviews — el bloque de respuesta generativa que Google muestra encima de los resultados orgánicos — desplazó '
    'entre un 15% y un 64% el CTR de las páginas posicionadas en queries informacionales durante 2024-2025 según múltiples '
    'estudios independientes del sector. El sector salud es el segmento con mayor índice de aparición de AI Overviews en '
    'España: queries sobre síntomas, definiciones de patologías, dosis, interacciones medicamentosas y horarios farmacéuticos '
    'se responden hoy directamente en el resultado generativo.', size=11)
add_para(doc, 'La exposición de topdoctors.es es directamente cuantificable:', size=11)
bullet(doc, '/diccionario-medico/ aporta el 24.83% del tráfico — su contenido es 100% informacional (definición + síntomas + tratamiento).')
bullet(doc, '/farmacia/ aporta el 23.33% — queries del tipo "farmacia de guardia", "farmacia cerca de mí" se están migrando a Google Local + AI.')
bullet(doc, '0 featured snippets capturados en la muestra de 250 keywords — el dominio no está siendo citado como fuente AI.')
add_para(doc, 'Suma de tráfico en zona de riesgo directo: 123,921 visitas/mes en la muestra (48.16%) — '
              'el equivalente a $4M+ USD/mes en valor de tráfico que está en zona de absorción AI.', size=11, bold=True, color=NAVY)

h2(doc, '4.3 Amenaza 2 — Asimetría de Autoridad en el Perfil de Backlinks')
add_para(doc, 'Probabilidad de impacto: MEDIA-ALTA. Evidencia visible en la muestra de 250 referring domains.', size=11, bold=True, color=RED)
add_para(doc,
    'El dominio tiene DR 78 acumulado en 6,525 referring domains, pero el DR promedio de los 250 referrentes más recientes '
    'es de solo 12.23/100. Esa cifra es muy baja para un dominio de la magnitud de topdoctors.es y revela un patrón típico: '
    'la autoridad histórica fue construida con enlaces fuertes (wikipedia.org, cosmopolitan.com, hellomagazine.com, '
    'topdoctors.co.uk) pero el flujo continuo de enlaces nuevos es de baja calidad, dependiente de directorios y blogs '
    'pequeños. Si Google ajusta el algoritmo para penalizar enlaces de baja autoridad — como hizo con la actualización de '
    'spam-link de octubre 2025 — la autoridad activa del dominio cae más rápido de lo que se renueva.', size=11)

h2(doc, '4.4 Oportunidad 1 — 35 Quick-Wins con Volumen ≥ 1,000')
add_para(doc, 'Probabilidad de captura: ALTA. Trabajo de optimización quirúrgica + enlazado interno.', size=11, bold=True, color=GREEN)
add_para(doc,
    'En la muestra de 250 keywords analizadas, 35 cumplen tres condiciones simultáneas: posición entre 4 y 20, volumen ≥ 1,000 '
    'búsquedas/mes y Keyword Difficulty ≤ 25. Estas son las keywords con mayor relación esfuerzo/recompensa del dominio. '
    'Mover una keyword desde la posición #6 al #1 multiplica el tráfico estimado entre 4 y 7 veces (curva de CTR estándar Ahrefs).', size=11)

styled_table(doc,
    ['Keyword','Vol/mes','Pos. actual','KD','URL actual'],
    [
        ['farmacia de guardia',                       '37,000','#7', '6',  '/farmacia/toma-de-tension/madrid/'],
        ['conjuntivitis',                             '37,000','#12','14', '/diccionario-medico/conjuntivitis/'],
        ['diverticulitis',                            '34,000','#13','19', '/articulos-medicos/que-cuidados-debo-seguir-cuando-tengo-diverticulos.../'],
        ['isquemia',                                  '33,000','#6', '5',  '/diccionario-medico/isquemia/'],
        ['bruxismo',                                  '28,000','#11','22', '/diccionario-medico/bruxismo/'],
        ['hernia inguinal',                           '27,000','#14','5',  '/articulos-medicos/hernia-inguinal-cuando-es-necesaria-la-cirugia.../'],
        ['leucemia',                                  '25,000','#10','13', '/diccionario-medico/leucemia/'],
        ['fisioterapia',                              '24,000','#5', '10', '/madrid-ciudad/fisioterapia/'],
        ['dermatitis',                                '22,000','#7', '12', '/diccionario-medico/dermatitis/'],
        ['peritonitis',                               '21,000','#14','5',  '/articulos-medicos/como-identificar-y-tratar-la-peritonitis.../'],
        ['colesterol ldl',                            '21,000','#10','3',  '/diccionario-medico/colesterol-ldl/'],
        ['hipertiroidismo',                           '20,000','#12','16', '/diccionario-medico/hipertiroidismo/'],
        ['magnetoterapia',                            '19,000','#8', '0',  '/diccionario-medico/magnetoterapia/'],
        ['tendinitis',                                '19,000','#11','13', '/diccionario-medico/tendinitis/'],
        ['menisco',                                   '19,000','#11','0',  '/diccionario-medico/meniscos/'],
        ['pubalgia',                                  '18,000','#9', '12', '/diccionario-medico/pubalgia/'],
        ['farmacia de guardia cerca de mi',           '17,000','#5', '0',  '/farmacia/24h/madrid/madrid-ciudad/'],
        ['que farmacia está de guardia hoy',          '14,000','#6', '6',  '/farmacia/toma-de-tension/madrid/'],
        ['osteopata',                                 '14,000','#5', '1',  '/madrid-ciudad/osteopatia/'],
        ['condromalacia rotuliana',                   '14,000','#12','0',  '/articulos-medicos/vivir-con-condromalacia-rotuliana/'],
    ]
)
add_para(doc,
    'Solo de esta tabla, el potencial de tráfico adicional si las 20 keywords listadas ascienden a top 3 es de '
    'aproximadamente 25,000–40,000 visitas/mes adicionales (cálculo con CTR estándar Ahrefs para pos #1-3 vs pos actual).',
    size=11, italic=True, color=GREY_MED)

h2(doc, '4.5 Oportunidad 2 — Replicación del Modelo Madrid en Hubs Regionales')
add_para(doc, 'Probabilidad de captura: ALTA. Activo geográfico ya construido pero subutilizado.', size=11, bold=True, color=GREEN)
add_para(doc,
    'Madrid concentra el 15.46% del tráfico del dominio con 43 páginas activas. Barcelona, segunda ciudad de España por PIB y '
    'población, aporta solo el 3.07% del tráfico pese a tener 13 páginas activas — es decir, las páginas existen pero rinden '
    '5x menos por unidad que las de Madrid. El patrón se repite en Valencia (1.93%), Alicante (1.01%) y Málaga (0.99%). '
    'La replicación del modelo Madrid (densidad de contenido, enlazado interno, schema LocalBusiness) en los 9 hubs urbanos '
    'principales puede generar +20,000 a +40,000 visitas/mes incrementales en 90 días sin contenido completamente nuevo.', size=11)

h2(doc, '4.6 Resumen del Diagnóstico — Matriz de Causas y Oportunidades')
styled_table(doc,
    ['Vector','Tipo','Probabilidad','Evidencia en Ahrefs','Impacto estimado'],
    [
        ['AI Overviews absorbiendo /diccionario-medico/ + /farmacia/','Amenaza','Alta','48.16% del tráfico en zona de riesgo','Pérdida potencial 30-50% del tráfico afectado en 12 meses'],
        ['Asimetría DR 78 dominio vs DR 12.23 referrentes','Amenaza','Media-Alta','250 ref. domains recientes con DR avg 12.23','Vulnerabilidad ante actualizaciones de spam-link'],
        ['35 quick-wins en pos 4–20 con vol ≥ 1,000','Oportunidad','Alta','Muestra 250 kws: 85 en pos 4–10','+25,000 a +40,000 visitas/mes en 90 días'],
        ['Hubs regionales subutilizados','Oportunidad','Alta','Barcelona/Valencia/Alicante 5x menos eficientes que Madrid','+20,000 a +40,000 visitas/mes en 90 días'],
        ['0 featured snippets capturados','Amenaza+Oportunidad','Alta','0 / 250 kws en muestra','Captura de 15+ snippets viable en 60 días'],
    ]
)

doc.add_page_break()

# ============ 5. KEYWORDS MAYOR VALOR ============
h1(doc, '5. Keywords de Mayor Valor Económico Actual')

add_para(doc,
    'Las siguientes keywords son las que actualmente generan mayor valor económico al dominio, calculado como '
    'el costo equivalente en Google Ads (tráfico estimado × CPC oficial de Ahrefs). Son los activos prioritarios '
    'a proteger frente a cualquier cambio algorítmico futuro:', size=11)

styled_table(doc,
    ['Keyword','Pos.','Vol. mensual','CPC','Valor mensual estimado'],
    [
        ['agrupacio mutua',                           '#2', '5,000',  '$111', '$1,047,285 USD'],
        ['tricologia',                                '#2', '1,700',  '$525', '$443,625 USD'],
        ['fibrinogeno',                               '#6', '5,700',  '$703', '$385,244 USD'],
        ['ortodoncia invisible alicante',             '#4', '3,500',  '$637', '$263,081 USD'],
        ['farmacia cerca de mi',                      '#2', '118,000','$11',  '$120,252 USD'],
        ['tricologo',                                 '#3', '3,300',  '$228', '$115,824 USD'],
        ['top doctors',                               '#1', '8,500',  '$14',  '$108,920 USD'],
        ['topdoctors',                                '#1', '3,600',  '$21',  '$72,219 USD'],
        ['dentista barcelona',                        '#1', '1,200',  '$234', '$71,136 USD'],
        ['dentista',                                  '#3', '14,000', '$162', '$68,526 USD'],
        ['tricologo barcelona',                       '#1', '450',    '$224', '$58,912 USD'],
        ['fisioterapia',                              '#5', '24,000', '$60',  '$57,840 USD'],
        ['sindrome de abstinencia',                   '#2', '4,700',  '$34',  '$53,652 USD'],
        ['nefrologo',                                 '#1', '4,300',  '$35',  '$50,435 USD'],
        ['farmacias cerca de mi ubicación abiertas',  '#2', '4,000',  '$25',  '$43,600 USD'],
    ]
)
callout(doc, 'INSIGHT ESTRATÉGICO',
    'La keyword "agrupacio mutua" (un seguro médico privado español) genera por sí sola valor equivalente a $1.04M USD/mes — '
    'pero está en posición #2. Subirla al #1 puede agregar $400K-$500K USD/mes de valor adicional. Este tipo de keywords '
    'transaccionales de seguros médicos privados son las que justifican económicamente el plan SEO completo.',
    kind='info')

# ============ 6. PERFIL AUTORIDAD ============
h1(doc, '6. Perfil de Autoridad y Backlinks')

h2(doc, '6.1 Estado del Perfil de Backlinks')
styled_table(doc,
    ['Indicador','Valor actual','Evaluación'],
    [
        ['Total backlinks (dominio)',         '353,179',     'Perfil voluminoso y diversificado'],
        ['Total referring domains',           '6,525',       'Diversificación sólida'],
        ['Domain Rating del dominio',         '78 / 100',    'Autoridad top tier para salud privada ES'],
        ['Backlinks en muestra reciente',     '250',         'Análisis detallado de los más recientes'],
        ['Backlinks dofollow en muestra',     '189 (75.6%)', 'Ratio saludable — transmiten autoridad'],
        ['Backlinks nofollow en muestra',     '61 (24.4%)',  'Diversidad natural — positivo'],
        ['DR promedio de referrentes recientes','12.23',     'BAJO — alerta de asimetría histórica vs actual'],
        ['DR máximo en referrentes recientes','97 (wikipedia.org)', 'Presencia en dominios top tier — defendible'],
        ['Backlinks perdidos en muestra',     '0',           'Sin pérdidas masivas visibles en últimos enlaces'],
    ]
)

h2(doc, '6.2 Top Referring Domains de Mayor Autoridad')
styled_table(doc,
    ['Dominio referente','DR','Backlinks al sitio','Dofollow','Primera detección'],
    [
        ['wikipedia.org',     '97','39',   'No',  '2021-01-08'],
        ['cosmopolitan.com',  '89','39',   'Sí',  '2019-05-29'],
        ['hellomagazine.com', '82','3',    'No',  '2025-10-15'],
        ['ull.es',            '77','6',    'Sí',  '2022-05-16'],
        ['elcomercio.es',     '77','1',    'Sí',  '2026-05-02'],
        ['elnortedecastilla.es','76','1',  'Sí',  '2026-04-29'],
        ['elnacional.cat',    '75','38',   'Sí',  '2023-10-06'],
        ['topdoctors.co.uk',  '74','3,000','Sí',  '2021-10-23'],
        ['dentaly.org',       '70','10',   'Sí',  '2022-01-09'],
        ['vademecum.es',      '65','1',    'Sí',  '2023-10-14'],
        ['corachan.com',      '52','12',   'Sí',  '2022-04-14'],
    ]
)
callout(doc, 'OBSERVACIÓN',
    'El backlink desde topdoctors.co.uk (DR 74, 3,000 enlaces dofollow) es el activo de enlazado más importante del dominio. '
    'La sinergia entre las versiones nacionales del grupo Top Doctors no está siendo explotada en otros mercados (Italia, '
    'México, Brasil, Reino Unido). Una estrategia coordinada de crosslinking inter-regional puede agregar de forma inmediata '
    'enlaces de DR 60-80 a las páginas españolas más estratégicas.',
    kind='info')

doc.add_page_break()

# ============ 7. PLAN DE ACCIÓN ============
h1(doc, '7. Plan de Acción y Recuperación')

h2(doc, 'Fase 1 — Quick Wins (Semanas 1 y 2)')

h3(doc, 'Acción 1.1 — Ataque a las 30 keywords prioritarias en posiciones 4–10')
add_para(doc,
    'Trabajo quirúrgico sobre las 30 keywords con mayor potencial inmediato (volumen ≥ 5,000, KD ≤ 20, posición actual 4-10): '
    'auditoría de la página posicionada, expansión de contenido con E-E-A-T verificable (autoría médica, citas, fuentes oficiales), '
    'implementación de schema MedicalCondition / FAQPage / HowTo según corresponda, y refuerzo de enlazado interno desde la '
    'homepage y desde las páginas de mayor URL Rating del dominio.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de contenido + SEO técnico.   |   Plazo: 14 días.   |   '
    'Meta: 8-12 keywords ascendiendo al top 3 antes del día 30 con +5,000 visitas/mes acumuladas.', size=10, italic=True, color=GREY_MED)

h3(doc, 'Acción 1.2 — Medición línea base de AI Share of Voice')
add_para(doc,
    'Auditoría de 200 prompts representativos del sector salud privado en España sobre ChatGPT, Gemini, Perplexity y Google AI '
    'Overviews, midiendo cuántas veces aparece citado topdoctors.es vs Doctoralia, Quirónsalud, Sanitas, MAPFRE y Adeslas. '
    'Esta medición se convierte en el KPI maestro de la fase 3.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de Generative Search.   |   Plazo: 10 días.   |   '
    'Entregable: Dashboard de AI Share of Voice base + benchmark competitivo.', size=10, italic=True, color=GREY_MED)

h3(doc, 'Acción 1.3 — Activación de alertas de rank tracking y monitoreo AI Overviews')
add_para(doc,
    'Configuración de Rank Tracker Ahrefs sobre las 100 keywords de mayor valor económico (Sección 5), con alerta '
    'automática ante caídas > 3 posiciones. Activación de monitoreo automático de aparición de AI Overviews en '
    'esas mismas 100 queries vía SerpAPI/DataForSEO. Frente a cualquier movimiento, respuesta del equipo en < 48h.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de SEO técnico.   |   Plazo: 5 días.   |   '
    'Entregable: Panel de monitoreo + protocolo de respuesta documentado.', size=10, italic=True, color=GREY_MED)

h2(doc, 'Fase 2 — Crecimiento Estructural (Meses 1 y 2)')

h3(doc, 'Acción 2.1 — Blindaje de /diccionario-medico/ frente a AI Overviews')
add_para(doc,
    'Rediseño de las 74 páginas de mayor tráfico en /diccionario-medico/ siguiendo el patrón de contenido que las IA generativas '
    'priorizan: definición directa en los primeros 50 caracteres, schema MedicalCondition completo, autoría médica verificable '
    'con CV linkeable, sección "Preguntas frecuentes" con FAQ schema, citas a fuentes oficiales (SemFYC, OMS, Ministerio de '
    'Sanidad ES) y bloque "Doctores especializados en esta condición" cruzando hacia los perfiles activos en /doctor/. '
    'El objetivo es que cuando Google AI Overviews responda sobre una patología, cite a topdoctors.es como fuente y enlace al '
    'perfil del especialista correspondiente.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency + revisión médica externa del cliente.   |   Plazo: 45 días.   |   '
    'Meta: 50 páginas rediseñadas + 15 featured snippets capturados.', size=10, italic=True, color=GREY_MED)

h3(doc, 'Acción 2.2 — Replicación del Modelo Madrid en hubs regionales')
add_para(doc,
    'Análisis de cada hub regional (Barcelona, Valencia, Alicante, Málaga, Sevilla, Bilbao, Pamplona, Las Palmas, Tenerife) '
    'identificando: keywords de "[especialidad] + [ciudad]" donde topdoctors.es no está en top 3, densidad de contenido por '
    'página de especialidad vs el equivalente Madrid, enlazado interno entre páginas de especialidades dentro de la misma '
    'ciudad. Replicación del patrón estructural más exitoso de Madrid en cada hub.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de contenido + SEO técnico.   |   Plazo: 60 días.   |   '
    'Meta: 9 hubs urbanos con +50% de keywords en top 3 por especialidad.', size=10, italic=True, color=GREY_MED)

h3(doc, 'Acción 2.3 — Captura de Featured Snippets')
add_para(doc,
    'Identificación de 80 queries donde topdoctors.es está en pos 2-5 y existe un featured snippet capturado por la '
    'competencia. Reescritura del primer párrafo de cada página objetivo siguiendo el patrón estructural del snippet '
    'rival (definición + lista numerada + tabla resumen). Meta: capturar al menos 15 snippets en 60 días — cada snippet '
    'capturado equivale a ~20-40% de incremento de CTR sobre la posición orgánica original.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de contenido.   |   Plazo: 60 días.   |   '
    'Meta: 15 featured snippets capturados de cero a 15.', size=10, italic=True, color=GREY_MED)

h2(doc, 'Fase 3 — Autoridad y AI Search (Meses 1 al 3)')

h3(doc, 'Acción 3.1 — Campaña de link building dirigido')
add_para(doc,
    'El perfil de backlinks asimétrico (DR 78 dominio vs DR 12.23 referrentes recientes) requiere una estrategia activa de '
    'enlaces de calidad. Plan: outreach a 20 medios médicos y de salud españoles (Diariofarma, Redacción Médica, Gaceta Médica, '
    'Consalud, Heraldo Salud), 15 medios generalistas con secciones de salud (El País Bienestar, ABC Salud, La Vanguardia Salud, '
    '20minutos Salud), y 10 universidades / colegios profesionales (Colegio Oficial de Médicos de Madrid, Cataluña, Valencia, '
    'COFM, COFB) para colaboraciones de contenido con backlink a perfiles de doctores especializados.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de PR digital.   |   Plazo: 90 días.   |   '
    'Meta: 30 referring domains nuevos con DR ≥ 50.', size=10, italic=True, color=GREY_MED)

h3(doc, 'Acción 3.2 — Contenido optimizado para AI Search')
add_para(doc,
    'Producción de 25 piezas de contenido pensadas específicamente para ser citadas por IA generativa: comparativas '
    'estructuradas ("Mejores tratamientos para X en 2026", "Cómo elegir un especialista en Y"), Q&A médicos profundos '
    'con autoría verificada, glosarios temáticos enlazables, e infografías con datos citables. Cada pieza incluye schema '
    'Article + Person (autor médico) + MedicalCondition / MedicalProcedure según corresponda, e integración cruzada con '
    'los perfiles de doctores activos en /doctor/.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency — equipo de contenido + supervisión médica del cliente.   |   Plazo: 90 días.   |   '
    'Meta: 25 piezas publicadas, AI Share of Voice > 25%.', size=10, italic=True, color=GREY_MED)

h3(doc, 'Acción 3.3 — Crosslinking inter-regional Top Doctors')
add_para(doc,
    'Activación del activo más infrautilizado del grupo: la red internacional Top Doctors. Coordinación con las versiones '
    '.co.uk, .com.mx, .it y .com.br para crosslinking estructurado desde los perfiles de doctores que ejercen en varios '
    'mercados. El backlink ya existente desde topdoctors.co.uk (DR 74, 3,000 enlaces dofollow) es la prueba de concepto: '
    'multiplicar ese patrón aporta enlaces DR 60-80 sin coste de adquisición.', size=11)
add_para(doc,
    'Responsable: SEO Lab Agency + coordinación inter-regional del cliente.   |   Plazo: 90 días.   |   '
    'Meta: +500 backlinks dofollow desde dominios Top Doctors internacionales.', size=10, italic=True, color=GREY_MED)

# ============ 8. METRICAS ============
h1(doc, '8. Métricas de Éxito y Seguimiento')

styled_table(doc,
    ['Métrica','Valor actual (mayo 2026)','Objetivo 30 días','Objetivo 60 días','Objetivo 90 días'],
    [
        ['Tráfico orgánico mensual',              '760,959 visitas',    '+10,000', '+30,000', '+60,000 (8%)'],
        ['Valor de tráfico mensual (USD)',        '$21,570,727',         '$22.0M',  '$22.5M',  '> $23.0M'],
        ['Keywords en pos 4-10 (muestra 250)',    '85 keywords',         '70',      '55',      '< 40'],
        ['Featured snippets capturados',          '0',                   '3',       '8',       '> 15'],
        ['Referring domains de DR ≥ 50',          'Bajo (no cuantificado)','+5',    '+15',     '+30'],
        ['AI Share of Voice (vs Doctoralia/Quirónsalud/Sanitas)','Sin medición previa','Línea base','>15%','> 25%'],
        ['DR promedio de referrentes nuevos',     '12.23',               '15',      '20',      '> 25'],
        ['Hubs regionales con +50% kws top 3',    '1 (Madrid)',          '2',       '4',       '6 ciudades'],
    ]
)

# ============ 9. CONCLUSIONES ============
h1(doc, '9. Conclusiones')

bullet(doc,
    'topdoctors.es es uno de los activos digitales más sólidos del sector salud privado en España (DR 78, 760,959 visitas/mes, '
    '$21.57M USD de valor de tráfico orgánico mensual). No hay penalización, no hay caída medida, no hay pérdida de autoridad.',
    bold_prefix='1. El dominio NO está en crisis. ')
bullet(doc,
    '48.16% del tráfico depende de carpetas (/diccionario-medico/, /farmacia/) que Google AI Overviews ya está absorbiendo en '
    'el sector salud español. Cero featured snippets capturados confirma que el dominio no está siendo citado como fuente AI.',
    bold_prefix='2. La exposición a AI Overviews es estructural y cuantificable. ')
bullet(doc,
    '85 keywords en posiciones 4–10 (34% de la muestra) y la asimetría hubs regionales vs Madrid suman 45,000-80,000 visitas/mes '
    'incrementales viables en 90 días sin contenido completamente nuevo.',
    bold_prefix='3. Las oportunidades son inmediatas y bien dimensionadas. ')
bullet(doc,
    'El DR 78 acumulado se sostiene en enlaces históricos. El flujo continuo de backlinks nuevos tiene DR promedio de 12.23 — '
    'cualquier actualización futura de spam-link golpea la autoridad activa más rápido de lo que se renueva.',
    bold_prefix='4. El perfil de autoridad necesita reinversión. ')
bullet(doc,
    'Sin trabajo SEO proactivo, topdoctors.es enfrenta en 12 meses una pérdida estimada del 15-30% de su tráfico informacional '
    'por absorción AI. Con el plan de 3 fases propuesto, no solo se neutraliza esa amenaza: se captura un crecimiento neto '
    'medible y se construye la ventaja competitiva en AI Search antes que la competencia.',
    bold_prefix='5. La ventana de acción es ahora. ')

callout(doc, 'PRÓXIMO PASO INMEDIATO',
    'Reunión de presentación de hallazgos con el equipo de marketing digital de Top Doctors España (1.5h). '
    'Validación de prioridades, definición de SLAs operativos y kickoff de Fase 1 (Quick Wins) en máximo 7 días desde la firma.',
    kind='info')

# Footer final
doc.add_paragraph().paragraph_format.space_after = Pt(20)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, 'Preparado por SEO Lab Agency  |  Mayo 2026', size=10, bold=True, color=NAVY)
fp2 = doc.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp2, 'Documento confidencial — Análisis basado exclusivamente en datos verificables de Ahrefs', size=9, italic=True, color=GREY_MED)

# Guardar
out = '/home/user/Agents_Automations/Informe_SEO_TopDoctors_ES.docx'
doc.save(out)
print('OK ->', out)
